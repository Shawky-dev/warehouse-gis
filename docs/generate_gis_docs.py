"""
Generate docs/warehouse_gis_technical_documentation.docx

Run from the repo root:
    python docs/generate_gis_docs.py
"""

import io
import os

import matplotlib
import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

matplotlib.use("Agg")

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY = RGBColor(0x1F, 0x4E, 0x79)        # #1F4E79
BLUE = RGBColor(0x2E, 0x75, 0xB6)        # #2E75B6
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_BG = RGBColor(0xF2, 0xF2, 0xF2)    # #F2F2F2

NAVY_HEX = "#1F4E79"
BLUE_HEX = "#2E75B6"
LIGHT_BLUE_HEX = "#D6E4F0"
RED_HEX = "#C00000"
AMBER_HEX = "#ED7D31"
GREEN_HEX = "#375623"
GRAY_HEX = "#808080"


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table-cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def add_paragraph_border(para, color_hex: str = "2E75B6", width_pt: int = 12):
    """Add a left border to a paragraph (used for code blocks)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_pt))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(left)
    pPr.append(pBdr)


def add_paragraph_shading(para, hex_color: str = "F2F2F2"):
    """Fill paragraph background (for code blocks)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    pPr.append(shd)


def set_para_spacing(para, before_pt: int = 0, after_pt: int = 0, line_spacing_pt: int = 0):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    if before_pt:
        spacing.set(qn("w:before"), str(before_pt * 20))
    if after_pt:
        spacing.set(qn("w:after"), str(after_pt * 20))
    if line_spacing_pt:
        spacing.set(qn("w:line"), str(line_spacing_pt * 20))
        spacing.set(qn("w:lineRule"), "exact")
    pPr.append(spacing)


def add_h1(doc, text: str) -> None:
    """Section heading — Calibri 18pt Bold, #1F4E79, bottom border."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(para, before_pt=18, after_pt=6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    # Bottom border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")     # 2pt = 16 eighths-of-a-point
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_h2(doc, text: str) -> None:
    """Sub-heading — Calibri 14pt Bold, #2E75B6."""
    para = doc.add_paragraph()
    set_para_spacing(para, before_pt=12, after_pt=4)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE


def add_h3(doc, text: str) -> None:
    """Sub-sub-heading — Calibri 12pt Bold, #1F4E79."""
    para = doc.add_paragraph()
    set_para_spacing(para, before_pt=8, after_pt=2)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY


def add_body(doc, text: str) -> None:
    """Normal body paragraph — Calibri 11pt, black."""
    para = doc.add_paragraph()
    set_para_spacing(para, before_pt=0, after_pt=6)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK


def add_code(doc, text: str) -> None:
    """Code block — Courier New 9pt, #F2F2F2 background, #2E75B6 left border."""
    for line in text.split("\n"):
        para = doc.add_paragraph()
        set_para_spacing(para, before_pt=0, after_pt=0, line_spacing_pt=12)
        add_paragraph_shading(para, "F2F2F2")
        add_paragraph_border(para, "2E75B6", width_pt=12)
        run = para.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = BLACK
    # small gap after block
    gap = doc.add_paragraph()
    set_para_spacing(gap, before_pt=0, after_pt=4)


def add_horizontal_rule(doc) -> None:
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)


def embed_figure(doc, fig, width_cm: float = 16.0) -> None:
    """Render a matplotlib figure to a PNG and embed it in the document."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(buf, width=Cm(width_cm))
    set_para_spacing(para, before_pt=6, after_pt=10)


# ── Diagram helpers ────────────────────────────────────────────────────────────

def make_box(ax, x, y, w, h, label, fontsize=8,
             facecolor=LIGHT_BLUE_HEX, edgecolor=NAVY_HEX, text_color="black",
             bold=False):
    rect = mpatches.FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.02",
        facecolor=facecolor,
        edgecolor=edgecolor,
        linewidth=1.2,
        zorder=2,
    )
    ax.add_patch(rect)
    weight = "bold" if bold else "normal"
    ax.text(x + w / 2, y + h / 2, label,
            ha="center", va="center", fontsize=fontsize,
            color=text_color, fontweight=weight, zorder=3,
            wrap=True)


def arrow(ax, x1, y1, x2, y2, color=GRAY_HEX):
    ax.annotate(
        "", xy=(x2, y2), xytext=(x1, y1),
        arrowprops=dict(
            arrowstyle="-|>",
            color=color,
            lw=1.4,
            mutation_scale=12,
        ),
        zorder=1,
    )


# ═══════════════════════════════════════════════════════════════════════════════
# Diagram 1 — System Architecture (swim-lane style)
# ═══════════════════════════════════════════════════════════════════════════════

def build_architecture_diagram():
    fig, ax = plt.subplots(figsize=(16, 10))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 10)
    ax.axis("off")

    # Column headers
    col_headers = [
        (0.3, "Browser\n(React + ArcGIS JS SDK 4)"),
        (5.6, "Spring Boot API"),
        (11.0, "Data / Infrastructure"),
    ]
    for cx, label in col_headers:
        rect = mpatches.FancyBboxPatch(
            (cx, 9.1), 4.8, 0.75,
            boxstyle="round,pad=0.04",
            facecolor=NAVY_HEX,
            edgecolor=NAVY_HEX,
            linewidth=0,
            zorder=2,
        )
        ax.add_patch(rect)
        ax.text(cx + 2.4, 9.47, label,
                ha="center", va="center", fontsize=9.5,
                color="white", fontweight="bold", zorder=3)

    # ── Column 1: Browser ─────────────────────────────────────────────────────
    browser_items = [
        "2D MapView",
        "SVG Floor Plan\n(MediaLayer)",
        "Zone Polygons\n(GraphicsLayer)",
        "Hazard Buffer Polygons\n(GraphicsLayer)",
        "Data Layers\n(MediaLayer)",
        "Click-to-Inspect Panel",
    ]
    bx, bw, bh, gap = 0.3, 4.8, 0.9, 0.22
    for i, label in enumerate(browser_items):
        by = 8.0 - i * (bh + gap)
        make_box(ax, bx, by, bw, bh, label, fontsize=8.2)

    # ── Column 2: Spring Boot ─────────────────────────────────────────────────
    api_items = [
        "GisLayerController\n(GeoJSON)",
        "GisZoneController\n(CRUD + GeoJSON)",
        "GisHazardBufferController\n(import)",
        "GisDataLayerController\n(upload / serve)",
        "FloorPlanController\n(manual drawing)",
        "InventoryValidationService",
        "GeoServerProvisioningService",
    ]
    sx, sw = 5.6, 4.8
    sh, sgap = 0.77, 0.18
    for i, label in enumerate(api_items):
        sy = 8.1 - i * (sh + sgap)
        make_box(ax, sx, sy, sw, sh, label, fontsize=7.8)

    # ── Column 3: Infrastructure ──────────────────────────────────────────────
    ix = 11.0
    # PostGIS block
    make_box(ax, ix, 5.6, 4.8, 3.1, "", fontsize=8, facecolor="#EBF5FB", edgecolor=NAVY_HEX)
    ax.text(ix + 2.4, 8.5, "PostgreSQL + PostGIS", ha="center", va="center",
            fontsize=8.5, color=NAVY_HEX, fontweight="bold", zorder=3)
    for j, tbl in enumerate(["gis_blocks", "gis_zones", "gis_hazard_buffers", "gis_data_layers"]):
        make_box(ax, ix + 0.2, 7.9 - j * 0.65, 4.4, 0.52, tbl,
                 fontsize=7.8, facecolor=LIGHT_BLUE_HEX, edgecolor=BLUE_HEX)
    # GeoServer block
    make_box(ax, ix, 4.3, 4.8, 1.0,
             "GeoServer\n(WMS layers per tenant)",
             fontsize=8.5, facecolor="#EBF5FB", edgecolor=NAVY_HEX, bold=True)

    # ── Arrows ────────────────────────────────────────────────────────────────
    # Browser → API
    for i in range(6):
        by_mid = 8.0 - i * (0.9 + 0.22) + 0.45
        sy_mid = 8.1 - min(i, 4) * (0.77 + 0.18) + 0.385
        arrow(ax, bx + bw, by_mid, sx, sy_mid)

    # API → PostGIS
    for i in range(5):
        sy_mid = 8.1 - i * (0.77 + 0.18) + 0.385
        arrow(ax, sx + sw, sy_mid, ix, 7.9 - min(i, 3) * 0.65 + 0.26)

    # API → GeoServer
    arrow(ax, sx + sw, 8.1 - 6 * (0.77 + 0.18) + 0.385, ix, 4.8)

    fig.tight_layout()
    return fig


# ═══════════════════════════════════════════════════════════════════════════════
# Diagram 2 — Floor Plan Base Layer flow
# ═══════════════════════════════════════════════════════════════════════════════

def build_floor_plan_diagram():
    steps = [
        "CAD Software",
        "SVG Export",
        "Upload to\nSpring Boot",
        "Meters→Degrees\nConversion",
        "ArcGIS\nMediaLayer",
        "2D MapView",
    ]
    fig, ax = plt.subplots(figsize=(14, 2.6))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 2.6)
    ax.axis("off")

    bw, bh = 1.8, 1.4
    gap = 0.5
    total = len(steps) * bw + (len(steps) - 1) * gap
    start_x = (14 - total) / 2

    for i, label in enumerate(steps):
        bx = start_x + i * (bw + gap)
        by = 0.6
        make_box(ax, bx, by, bw, bh, label, fontsize=8.5, facecolor=LIGHT_BLUE_HEX)
        if i < len(steps) - 1:
            arrow(ax, bx + bw, by + bh / 2, bx + bw + gap, by + bh / 2)

    fig.tight_layout()
    return fig


# ═══════════════════════════════════════════════════════════════════════════════
# Diagram 3 — GIS Blocks vertical flow
# ═══════════════════════════════════════════════════════════════════════════════

def build_gis_blocks_diagram():
    steps = [
        "User draws polygon\non map",
        "FloorPlansPage\nEditor",
        "POST /blocks/manual",
        "FloorPlanController",
        "gis_blocks\n(PostGIS)",
        "GeoJSON Endpoint",
        "Map Viewer",
    ]
    fig, ax = plt.subplots(figsize=(4.5, 11))
    ax.set_xlim(0, 4.5)
    ax.set_ylim(0, 11)
    ax.axis("off")

    bw, bh = 3.4, 1.0
    gap = 0.35
    bx = 0.55
    total_h = len(steps) * bh + (len(steps) - 1) * gap
    start_y = (11 - total_h) / 2 + total_h

    for i, label in enumerate(steps):
        by = start_y - i * (bh + gap) - bh
        make_box(ax, bx, by, bw, bh, label, fontsize=8.5, facecolor=LIGHT_BLUE_HEX)
        if i < len(steps) - 1:
            cy = by
            arrow(ax, bx + bw / 2, cy, bx + bw / 2, cy - gap)

    fig.tight_layout()
    return fig


# ═══════════════════════════════════════════════════════════════════════════════
# Diagram 4 — Zone Management
# ═══════════════════════════════════════════════════════════════════════════════

def build_zones_diagram():
    fig, ax = plt.subplots(figsize=(14, 5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 5)
    ax.axis("off")

    # Zone entity box
    zone_label = (
        "GisZone\n"
        "─────────────────\n"
        "name\n"
        "geometry (Polygon/4326)\n"
        "violationAction: BLOCK | WARN\n"
        "categoryRules [ALLOWED|PROHIBITED]\n"
        "zoneType (FK)\n"
        "displayColor (#RRGGBB)\n"
        "source: MANUAL | ARCGIS_IMPORT"
    )
    make_box(ax, 4.5, 0.4, 5.0, 4.2, zone_label,
             fontsize=8.3, facecolor=LIGHT_BLUE_HEX, edgecolor=NAVY_HEX)

    # Manual Draw path
    make_box(ax, 0.2, 3.3, 3.5, 0.8, "Manual Draw on Map", fontsize=8.5)
    make_box(ax, 0.2, 2.3, 3.5, 0.8, "POST /gis/zones", fontsize=8.5)
    arrow(ax, 1.95, 3.3, 1.95, 3.1)
    arrow(ax, 1.95, 2.3, 1.95, 2.1)
    make_box(ax, 0.2, 1.3, 3.5, 0.8, "PostGIS gis_zones", fontsize=8.5)
    arrow(ax, 3.7, 2.65, 4.5, 2.5)

    # ArcGIS Import path
    make_box(ax, 10.3, 3.3, 3.5, 0.8, "ArcGIS Pro Export\n(GeoJSON)", fontsize=8.5)
    make_box(ax, 10.3, 2.3, 3.5, 0.8, "POST /zones/import", fontsize=8.5)
    arrow(ax, 12.05, 3.3, 12.05, 3.1)
    arrow(ax, 12.05, 2.3, 12.05, 2.1)
    make_box(ax, 10.3, 1.3, 3.5, 0.8, "PostGIS gis_zones", fontsize=8.5)
    arrow(ax, 10.3, 1.7, 9.5, 2.5)

    # Output
    make_box(ax, 4.5, 0.3, 5.0, 0.7, "GeoJSON → ArcGIS MapView", fontsize=8.5,
             facecolor=NAVY_HEX, edgecolor=NAVY_HEX, text_color="white")
    arrow(ax, 7.0, 0.4, 7.0, 0.3)

    fig.tight_layout()
    return fig


# ═══════════════════════════════════════════════════════════════════════════════
# Diagram 5 — Hazard Buffer flow
# ═══════════════════════════════════════════════════════════════════════════════

def build_hazard_buffer_diagram():
    steps = [
        "Safety Engineer",
        "ArcGIS Pro\nGeoJSON File",
        "Import API\n(/hazard-buffers/import)",
        "gis_hazard_buffers\n(PostGIS)",
        "ST_Intersects\nQuery",
        "InventoryValidation\nService",
        "[HARD BLOCK]",
    ]
    colors = [
        LIGHT_BLUE_HEX, LIGHT_BLUE_HEX, LIGHT_BLUE_HEX,
        LIGHT_BLUE_HEX, LIGHT_BLUE_HEX, LIGHT_BLUE_HEX,
        "#FFB3B3",
    ]
    fig, ax = plt.subplots(figsize=(15, 2.8))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 2.8)
    ax.axis("off")

    bw, bh = 1.8, 1.6
    gap = 0.4
    total = len(steps) * bw + (len(steps) - 1) * gap
    start_x = (15 - total) / 2

    for i, (label, fc) in enumerate(zip(steps, colors)):
        bx = start_x + i * (bw + gap)
        by = 0.6
        ec = RED_HEX if fc == "#FFB3B3" else NAVY_HEX
        tc = "black" if fc != "#FFB3B3" else "#8B0000"
        make_box(ax, bx, by, bw, bh, label, fontsize=8,
                 facecolor=fc, edgecolor=ec, text_color=tc)
        if i < len(steps) - 1:
            arrow(ax, bx + bw, by + bh / 2, bx + bw + gap, by + bh / 2)

    fig.tight_layout()
    return fig


# ═══════════════════════════════════════════════════════════════════════════════
# Diagram 6 — Spatial Rule Enforcement pipeline
# ═══════════════════════════════════════════════════════════════════════════════

def build_enforcement_diagram():
    fig, ax = plt.subplots(figsize=(9, 16))
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 16)
    ax.axis("off")

    cx = 4.5     # center-x
    bw = 6.5
    bx = (9 - bw) / 2

    # ── Top: trigger ──────────────────────────────────────────────────────────
    make_box(ax, bx, 14.5, bw, 1.0,
             "Product stored at location", fontsize=9,
             facecolor=NAVY_HEX, edgecolor=NAVY_HEX, text_color="white", bold=True)
    arrow(ax, cx, 14.5, cx, 13.9)

    # ── Step 1 ────────────────────────────────────────────────────────────────
    make_box(ax, bx, 12.5, bw, 1.3,
             "Step 1 — Hazard Buffer Check\n"
             "ST_Intersects(location_polygon, hazard_buffer)\n"
             "+ matching hazard type?",
             fontsize=8.5, facecolor="#EBF5FB", edgecolor=NAVY_HEX)

    # YES branch → BLOCK
    make_box(ax, bx + bw + 0.15, 12.6, 1.5, 0.8,
             "YES →", fontsize=8, facecolor="white", edgecolor="white")
    make_box(ax, bx + bw + 0.2, 11.5, 2.1, 0.9,
             "HARD BLOCK\n(never overrideable)",
             fontsize=8, facecolor="#FFB3B3", edgecolor=RED_HEX, text_color="#8B0000")
    arrow(ax, cx + bw / 2, 13.0, bx + bw + 0.8, 12.4)

    # NO → continue
    arrow(ax, cx, 12.5, cx, 11.9)
    ax.text(cx + 0.15, 12.2, "NO", fontsize=8, color=GRAY_HEX)

    # ── Step 2 ────────────────────────────────────────────────────────────────
    make_box(ax, bx, 10.4, bw, 1.4,
             "Step 2 — Zone Category Rule Check\n"
             "ST_Contains(zone_polygon, location_polygon)\n"
             "+ category PROHIBITED in zone?",
             fontsize=8.5, facecolor="#EBF5FB", edgecolor=NAVY_HEX)

    # YES BLOCK
    make_box(ax, bx + bw + 0.2, 10.8, 2.1, 0.7,
             "BLOCK\n(BLOCK zone)",
             fontsize=7.5, facecolor="#FFB3B3", edgecolor=RED_HEX, text_color="#8B0000")
    arrow(ax, cx + bw / 2, 11.0, bx + bw + 0.8, 11.1)

    # YES WARN
    make_box(ax, bx + bw + 0.2, 9.8, 2.1, 0.7,
             "WARN\n(WARN zone)",
             fontsize=7.5, facecolor="#FFF2CC", edgecolor=AMBER_HEX, text_color="#7F4C00")
    arrow(ax, cx + bw / 2, 10.8, bx + bw + 0.8, 10.1)

    # NO → continue
    arrow(ax, cx, 10.4, cx, 9.7)
    ax.text(cx + 0.15, 10.05, "NO / override=true", fontsize=8, color=GRAY_HEX)

    # ── Step 3 ────────────────────────────────────────────────────────────────
    make_box(ax, bx, 8.2, bw, 1.4,
             "Step 3 — Required Zone Type Check\n"
             "Product category requires zone type?\n"
             "ST_Contains(zone of required type, location)?",
             fontsize=8.5, facecolor="#EBF5FB", edgecolor=NAVY_HEX)

    # NO → WARN
    make_box(ax, bx + bw + 0.2, 8.5, 2.4, 0.8,
             "WARN + suggest\ncorrect zones",
             fontsize=7.5, facecolor="#FFF2CC", edgecolor=AMBER_HEX, text_color="#7F4C00")
    arrow(ax, cx + bw / 2, 8.8, bx + bw + 0.8, 8.85)
    ax.text(cx + bw / 2 + 0.1, 8.95, "NO", fontsize=8, color=GRAY_HEX)

    # YES → success
    arrow(ax, cx, 8.2, cx, 7.5)
    ax.text(cx + 0.15, 7.85, "YES", fontsize=8, color=GRAY_HEX)

    make_box(ax, bx, 6.5, bw, 0.9,
             "PASS — Stock movement recorded",
             fontsize=9, facecolor="#E2EFDA", edgecolor=GREEN_HEX, text_color="#375623", bold=True)

    fig.tight_layout()
    return fig


# ═══════════════════════════════════════════════════════════════════════════════
# Diagram 7 — Data Layers flow
# ═══════════════════════════════════════════════════════════════════════════════

def build_data_layers_diagram():
    steps = [
        "ArcGIS Pro\nAnalysis",
        "PNG / JPEG\nExport",
        "Upload via\nWeb UI",
        "Stored on\nDisk",
        "Served as\nbyte[]",
        "ArcGIS\nMediaLayer",
        "Opacity /\nOffset controls",
    ]
    fig, ax = plt.subplots(figsize=(15, 2.8))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 2.8)
    ax.axis("off")

    bw, bh = 1.8, 1.6
    gap = 0.4
    total = len(steps) * bw + (len(steps) - 1) * gap
    start_x = (15 - total) / 2

    for i, label in enumerate(steps):
        bx = start_x + i * (bw + gap)
        by = 0.6
        make_box(ax, bx, by, bw, bh, label, fontsize=8.5, facecolor=LIGHT_BLUE_HEX)
        if i < len(steps) - 1:
            arrow(ax, bx + bw, by + bh / 2, bx + bw + gap, by + bh / 2)

    fig.tight_layout()
    return fig


# ═══════════════════════════════════════════════════════════════════════════════
# Page setup helpers
# ═══════════════════════════════════════════════════════════════════════════════

def add_page_number_footer(doc):
    """Insert «Page X of Y» field in the footer of every section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.clear()

        run = para.add_run("Page ")
        run.font.name = "Calibri"
        run.font.size = Pt(9)

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.extend([fldChar1, instrText, fldChar2])

        run2 = para.add_run(" of ")
        run2.font.name = "Calibri"
        run2.font.size = Pt(9)

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "begin")
        instrText2 = OxmlElement("w:instrText")
        instrText2.text = "NUMPAGES"
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "end")
        run2._r.extend([fldChar3, instrText2, fldChar4])


# ═══════════════════════════════════════════════════════════════════════════════
# Styled table helpers
# ═══════════════════════════════════════════════════════════════════════════════

def add_data_model_table(doc):
    # Each entry: (table_name, key_columns, postgis_type, purpose)
    rows = [
        {
            "table":        "gis_blocks",
            "columns":      "layout_block_id, geometry, centroid_geom, depth",
            "postgis_type": "Polygon, Point (4326)",
            "purpose":      "Location polygons",
        },
        {
            "table":        "gis_zones",
            "columns":      "name, geometry, violation_action, zone_type_id",
            "postgis_type": "Polygon (4326)",
            "purpose":      "Storage zone polygons",
        },
        {
            "table":        "gis_hazard_buffers",
            "columns":      "name, geometry, import_batch_id",
            "postgis_type": "Polygon (4326)",
            "purpose":      "Danger zone polygons",
        },
        {
            "table":        "gis_data_layers",
            "columns":      "name, file_path, media_type",
            "postgis_type": "\u2014",
            "purpose":      "Heatmap image metadata",
        },
    ]
    field_order = ("table", "columns", "postgis_type", "purpose")
    headers = ["Table", "Key Columns", "PostGIS Type", "Purpose"]
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Table Grid"

    # Header row
    hrow = table.rows[0]
    for j, h in enumerate(headers):
        cell = hrow.cells[j]
        set_cell_bg(cell, "1F4E79")
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    # Data rows
    for i, row_dict in enumerate(rows):
        row = table.rows[i + 1]
        bg = "EBF5FB" if i % 2 == 0 else "FFFFFF"
        for j, field in enumerate(field_order):
            cell_text = row_dict[field]
            cell = row.cells[j]
            set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if j == 0:
                run.bold = True
    doc.add_paragraph()


def add_spatial_query_table(doc):
    rows = [
        ("Find buffers intersecting a location", "ST_Intersects", "GisHazardBufferRepository"),
        ("Find zones containing a location",     "ST_Contains",   "GisZoneRepository"),
        ("Check required zone type",             "ST_Contains",   "GisZoneRepository"),
        ("Leaf block identification",
         "Subquery (NOT IN parent_ids)",
         "GisBlockRepository"),
    ]
    headers = ["Operation", "PostGIS Function", "Used In"]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"

    hrow = table.rows[0]
    for j, h in enumerate(headers):
        cell = hrow.cells[j]
        set_cell_bg(cell, "1F4E79")
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    for i, row_data in enumerate(rows):
        row = table.rows[i + 1]
        bg = "EBF5FB" if i % 2 == 0 else "FFFFFF"
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# Main document builder
# ═══════════════════════════════════════════════════════════════════════════════

def build_document() -> Document:
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    add_page_number_footer(doc)

    # ── Cover Page ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Warehouse GIS Subsystem — Technical Documentation")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = NAVY

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        "Spatial Architecture, Feature Reference & Integration Guide"
    )
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = BLUE

    doc.add_paragraph()
    ver_para = doc.add_paragraph()
    ver_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_run = ver_para.add_run("v1.0  |  April 2026")
    ver_run.font.name = "Calibri"
    ver_run.font.size = Pt(11)
    ver_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_page_break()

    # ── Section 1 — Executive Summary ─────────────────────────────────────────
    add_h1(doc, "1 — Executive Summary")

    add_body(doc,
        "The Warehouse GIS subsystem transforms a conventional Warehouse Management System into a "
        "spatially-aware platform. Rather than treating warehouse locations as abstract codes in a "
        "database, the system anchors every physical location to a real-world coordinate system, "
        "enabling administrators to reason about storage rules in terms of geography — not just "
        "business logic. The result is a living floor plan that enforces storage constraints "
        "automatically as stock movements happen.")

    add_body(doc,
        "Administrators start by uploading an SVG floor plan exported from CAD software. They then "
        "draw polygons directly on a 2D map to associate each physical shelf, bay, or aisle with its "
        "logical counterpart in the WMS layout tree. On top of that spatial foundation they can define "
        "named zones — each carrying violation rules that either block or warn when incompatible "
        "products are placed inside them — and import hazard buffer polygons produced by safety "
        "engineers in ArcGIS Pro. Every time a stock movement is requested, the platform runs a "
        "three-step spatial pipeline against PostGIS to enforce those rules in real time.")

    add_body(doc,
        "The technology stack is designed for reliability and incremental adoption. The backend runs "
        "on Spring Boot with Hibernate Spatial exposing PostGIS geometry operations through JPA "
        "repositories. The frontend is a React single-page application that embeds the ArcGIS "
        "JavaScript SDK 4, giving users a responsive, interactive 2D map experience without requiring "
        "any GIS expertise. GeoServer sits alongside PostGIS to provide standards-based WMS layers "
        "for tenants who want to consume spatial data in external GIS tools such as ArcGIS Pro.")

    doc.add_page_break()

    # ── Section 2 — System Architecture ───────────────────────────────────────
    add_h1(doc, "2 — System Architecture")

    embed_figure(doc, build_architecture_diagram(), width_cm=16.5)

    add_body(doc,
        "The system follows a clean three-tier separation. The browser layer is built entirely with "
        "React components that embed an ArcGIS JS SDK 4 MapView. Spatial data arrives as standard "
        "GeoJSON from REST endpoints, which the SDK renders as GraphicsLayer or MediaLayer overlays "
        "on a 2D map. Interaction is intentionally lightweight: the SDK handles pan, zoom, and "
        "click-to-inspect while React manages application state and API calls.")

    add_body(doc,
        "Spring Boot sits in the middle, exposing a set of focused REST controllers — one each for "
        "floor plan management, GIS blocks, zones, hazard buffers, and data layers. All spatial "
        "persistence flows through PostGIS via JPA repositories that execute native SQL functions "
        "such as ST_Intersects and ST_Contains. GeoServer is provisioned automatically when new "
        "tenant workspaces are created, ensuring that WMS layers are always in sync with the "
        "underlying PostGIS tables.")

    doc.add_page_break()

    # ── Section 3 — Floor Plan Base Layer ─────────────────────────────────────
    add_h1(doc, "3 — Floor Plan Base Layer")

    add_h2(doc, "Data Flow")
    embed_figure(doc, build_floor_plan_diagram(), width_cm=15.0)

    add_body(doc,
        "The first step in making a warehouse spatially aware is placing its floor plan on a real "
        "coordinate system. Administrators upload an SVG floor plan exported from CAD software. "
        "The system converts physical dimensions (meters) into geographic coordinates (EPSG:4326) "
        "using the approximation that 1 degree ≈ 111 km. This is an equatorial approximation and "
        "works well for indoor/warehouse contexts where the coordinate span is small (typically "
        "< 0.01°), but administrators working at higher latitudes should be aware that longitude "
        "spans compress as latitude increases. The anchor is configured per-deployment and the "
        "SVG is anchored at a configurable lat/lon. "
        "The SVG is rendered as a georeferenced MediaLayer on the ArcGIS 2D map, giving the team an "
        "accurate visual backdrop before they draw any location polygons.")

    add_h3(doc, "Application Configuration (application.yml)")
    add_code(doc,
"""warehouse:
  gis:
    anchor-lat: 0.0
    anchor-lon: 0.0
    width-meters: 100.0
    length-meters: 60.0""")

    add_h3(doc, "TypeScript — Coordinate Conversion")
    add_code(doc,
"""const lonSpan = widthMeters / 111_000;
const latSpan = lengthMeters / 111_000;
const warehouseExtent = new Extent({
    xmin: anchorLon,  ymin: anchorLat,
    xmax: anchorLon + lonSpan,  ymax: anchorLat + latSpan,
    spatialReference: { wkid: 4326 },
});""")

    doc.add_page_break()

    # ── Section 4 — GIS Blocks ─────────────────────────────────────────────────
    add_h1(doc, "4 — GIS Blocks (Location Polygons)")

    add_h2(doc, "Data Flow")
    embed_figure(doc, build_gis_blocks_diagram(), width_cm=5.5)

    add_body(doc,
        "The logical warehouse layout tree — Warehouse → Zone → Aisle → Bay → Shelf — carries no "
        "spatial data by default. Administrators bridge this gap by opening the Floor Plans editor "
        "and drawing a polygon on the map for each physical location. Once drawn, the polygon is "
        "submitted to the backend which stores it as a GisBlock entity, linking a PostGIS Polygon "
        "geometry to the corresponding layoutBlockId. Locations that have no GIS block are treated "
        "gracefully: all spatial checks are simply skipped, so partially mapped warehouses continue "
        "to work without errors.")

    add_h3(doc, "GisBlock.java — JPA Entity")
    add_code(doc,
"""@Entity
@Table(name = "gis_blocks")
public class GisBlock {

    @Id
    private UUID id;

    @Column(name = "layout_block_id", nullable = false)
    private UUID layoutBlockId;

    @Column(name = "template_name", nullable = false, length = 100)
    private String templateName;

    @Column(length = 200)
    private String label;

    @Column(name = "position_path")
    private String positionPath;

    @Column(name = "depth", nullable = false)
    private int depth;

    @Column(columnDefinition = "GEOMETRY(Polygon, 4326)", nullable = false)
    private Polygon geometry;

    @Column(name = "centroid_geom", columnDefinition = "GEOMETRY(Point, 4326)")
    private Point centroidGeom;

    @Column(name = "created_at", nullable = false, updatable = false)
    private Instant createdAt;

    @Column(name = "updated_at", nullable = false)
    private Instant updatedAt;
}""")

    add_h3(doc, "FloorPlanController.java — saveManualBlock()")
    add_code(doc,
"""@PostMapping("/blocks/manual")
@PreAuthorize("hasAuthority(...GIS_FLOOR_PLAN_MANAGE)")
public ResponseEntity<?> saveManualBlock(
        @PathVariable String tenantSlug,
        @RequestBody ManualGisBlockRequest request,
        Authentication authentication) {

    tenantAccessPolicy.assertTenantAccess(authentication, tenantSlug);
    List<List<Double>> exteriorCoords = request.getRings().get(0);
    Coordinate[] coordinates = new Coordinate[exteriorCoords.size()];
    for (int i = 0; i < exteriorCoords.size(); i++) {
        List<Double> pt = exteriorCoords.get(i);
        coordinates[i] = new Coordinate(pt.get(0), pt.get(1));
    }
    GeometryFactory gf = new GeometryFactory(new PrecisionModel(), 4326);
    Polygon polygon = gf.createPolygon(gf.createLinearRing(coordinates));
    Point centroid = polygon.getCentroid();
    centroid.setSRID(4326);

    GisBlock gisBlock = gisBlockRepository
            .findByLayoutBlockId(request.getLayoutBlockId())
            .orElseGet(GisBlock::new);

    gisBlock.setLayoutBlockId(request.getLayoutBlockId());
    gisBlock.setGeometry(polygon);
    gisBlock.setCentroidGeom(centroid);
    gisBlock = gisBlockRepository.save(gisBlock);
    return ResponseEntity.ok(Map.of("id", gisBlock.getId()));
}""")

    add_h3(doc, "useEditorState.ts — savePolygon()")
    add_code(doc,
"""const savePolygon = useCallback(
  async (layoutBlockId: string, fullCode: string) => {
    if (!pendingPolygon) return;
    const depth = availableBlocks.find((b) => b.id === layoutBlockId)?.depth ?? 0;
    await api.post(`/${slug}/gis/blocks/manual`, {
      layoutBlockId,
      templateName: pendingPolygon.templateName,
      label: fullCode,
      positionPath: fullCode,
      depth,
      rings: pendingPolygon.rings,
    });
    await fetchExistingPolygons(pendingPolygon.templateName);
    setPendingPolygon(null);
  },
  [slug, pendingPolygon, availableBlocks, fetchExistingPolygons]
);""")

    doc.add_page_break()

    # ── Section 5 — GeoJSON Endpoints ─────────────────────────────────────────
    add_h1(doc, "5 — GeoJSON Endpoints")

    add_body(doc,
        "The frontend receives all spatial data as standard GeoJSON FeatureCollections from REST "
        "endpoints. This keeps the client-side code decoupled from the database schema: the SDK "
        "simply reads a FeatureCollection and renders it onto the map without knowing anything "
        "about the underlying JPA entities. Two primary endpoints feed the map viewer.")

    add_h3(doc, "Leaf Location Query (GisBlockRepository.java)")
    add_code(doc,
"""@Query(value = \"\"\"
    SELECT b.* FROM gis_blocks b
    WHERE b.layout_block_id NOT IN (
        SELECT DISTINCT lb.parent_id
        FROM layout_blocks lb
        WHERE lb.parent_id IS NOT NULL
    )
    \"\"\", nativeQuery = true)
List<GisBlock> findLeafGisBlocks();""")

    add_h3(doc, "Zones GeoJSON Endpoint (GisZoneController.java)")
    add_code(doc,
"""@GetMapping(value = "/geojson", produces = MediaType.APPLICATION_JSON_VALUE)
@PreAuthorize("hasAuthority(...GIS_ZONES_VIEW)")
@Transactional(readOnly = true)
public ResponseEntity<String> getZonesGeoJson(
        @PathVariable String tenantSlug,
        Authentication authentication) {
    tenantAccessPolicy.assertTenantAccess(authentication, tenantSlug);
    List<GisZone> zones = gisZoneRepository.findAllByOrderByCreatedAtAsc();
    return ResponseEntity.ok(buildZoneFeatureCollection(zones));
}""")

    doc.add_page_break()

    # ── Section 6 — Zone Management ───────────────────────────────────────────
    add_h1(doc, "6 — Zone Management")

    add_h2(doc, "Zone Lifecycle")
    embed_figure(doc, build_zones_diagram(), width_cm=15.0)

    add_body(doc,
        "Zones are named polygon areas that carry storage rules. Each zone has a violation action "
        "of either BLOCK or WARN: a BLOCK zone prevents the stock movement outright, while a WARN "
        "zone raises an advisory that operators can choose to override. Within each zone, category "
        "rules explicitly list which product categories are ALLOWED or PROHIBITED, giving "
        "administrators fine-grained control over what can be stored where.")

    add_body(doc,
        "Zones can be created in two ways. Administrators with the GIS_ZONES_MANAGE permission can "
        "draw a polygon directly on the map and submit it through the web UI, which calls the "
        "create endpoint. GIS analysts working in ArcGIS Pro can export a GeoJSON FeatureCollection "
        "and upload it through the import API, which batch-inserts all features in a single "
        "transaction. Both paths produce identical GisZone entities in PostGIS.")

    add_h3(doc, "GisZone.java — JPA Entity")
    add_code(doc,
"""@Entity
@Table(name = "gis_zones")
public class GisZone {

    @Id private UUID id;

    @Column(nullable = false, length = 200)
    private String name;

    @Column(columnDefinition = "GEOMETRY(Polygon, 4326)", nullable = false)
    private Polygon geometry;

    /** BLOCK | WARN */
    @Column(name = "violation_action", nullable = false, length = 20)
    private String violationAction;

    /** MANUAL | ARCGIS_IMPORT */
    @Column(nullable = false, length = 20)
    private String source;

    @ManyToOne(fetch = FetchType.LAZY)
    @JoinColumn(name = "zone_type_id")
    private ZoneType zoneType;

    @Column(name = "display_color", length = 7)
    private String displayColor;

    @OneToMany(mappedBy = "zone", cascade = CascadeType.ALL, orphanRemoval = true)
    private List<GisZoneCategoryRule> categoryRules = new ArrayList<>();
}""")

    add_h3(doc, "GisZoneController.java — createZone()")
    add_code(doc,
"""@PostMapping
@PreAuthorize("hasAuthority(...GIS_ZONES_MANAGE)")
@Transactional
public ResponseEntity<ZoneResponse> createZone(
        @PathVariable String tenantSlug,
        @Valid @RequestBody ZoneRequest request,
        Authentication authentication) {
    tenantAccessPolicy.assertTenantAccess(authentication, tenantSlug);

    GisZone zone = GisZone.builder()
            .name(request.name())
            .description(request.description())
            .geometry(ringsToPolygon(request.coordinates()))
            .violationAction(request.violationAction())
            .source(request.source() != null ? request.source() : "MANUAL")
            .zoneType(resolveZoneType(request.zoneTypeId()))
            .displayColor(normalizeDisplayColor(request.displayColor()))
            .build();

    GisZone saved = gisZoneRepository.save(zone);
    saved = applyRules(saved, request.categoryRules());
    return ResponseEntity.ok(toResponse(saved));
}""")

    add_h3(doc, "GisZoneController.java — importZones()")
    add_code(doc,
"""@PostMapping("/import")
@PreAuthorize("hasAuthority(...GIS_ZONES_MANAGE)")
@Transactional
public ResponseEntity<List<ZoneResponse>> importZones(
        @PathVariable String tenantSlug,
        @RequestBody GeoJsonImportRequest importRequest,
        Authentication authentication) {
    tenantAccessPolicy.assertTenantAccess(authentication, tenantSlug);

    List<ZoneResponse> created = importRequest.features().stream()
            .map(feature -> {
                GisZone zone = GisZone.builder()
                        .name(feature.properties().name())
                        .geometry(ringsToPolygon(extractRings(feature.geometry())))
                        .violationAction(feature.properties().violationAction() != null
                                ? feature.properties().violationAction() : "BLOCK")
                        .source("ARCGIS_IMPORT")
                        .build();
                return toResponse(gisZoneRepository.save(zone));
            })
            .toList();

    return ResponseEntity.ok(created);
}""")

    doc.add_page_break()

    # ── Section 7 — Hazard Buffers ─────────────────────────────────────────────
    add_h1(doc, "7 — Hazard Buffers")

    add_h2(doc, "Import & Enforcement Flow")
    embed_figure(doc, build_hazard_buffer_diagram(), width_cm=15.5)

    add_body(doc,
        "Hazard buffers are hard-block danger zones that mark areas of the warehouse where certain "
        "hazard types may never be stored. Unlike zones, which can be configured to warn instead of "
        "block, a hazard buffer violation is always a hard stop — there is no override mechanism. "
        "This makes them suitable for safety-critical boundaries such as blast radii, flammable "
        "material exclusion zones, and chemical incompatibility areas.")

    add_body(doc,
        "Buffers are always imported from ArcGIS Pro GeoJSON files rather than drawn manually. "
        "This ensures there is an auditable trail back to a certified safety analysis: each import "
        "batch receives a unique batchId UUID and the original filename is stored alongside the "
        "geometry. Safety engineers export a GeoJSON FeatureCollection from ArcGIS Pro and upload "
        "it through the web interface; Spring Boot parses the features, resolves the hazard type "
        "codes, and persists each buffer in one transaction.")

    add_h3(doc, "GisHazardBuffer.java — JPA Entity")
    add_code(doc,
"""@Entity
@Table(name = "gis_hazard_buffers")
public class GisHazardBuffer {

    @Id private UUID id;

    @Column(nullable = false, length = 200)
    private String name;

    @Column(nullable = false, length = 20)
    private String source;   // always "ARCGIS_IMPORT"

    @Column(columnDefinition = "GEOMETRY(Polygon, 4326)", nullable = false)
    private Polygon geometry;

    @Column(columnDefinition = "TEXT")
    private String notes;

    @Column(name = "import_batch_id")
    private UUID importBatchId;

    @Column(name = "source_filename", length = 255)
    private String sourceFilename;

    @Column(name = "imported_at", nullable = false)
    private Instant importedAt;

    @ManyToMany
    @JoinTable(name = "gis_hazard_buffer_restricted_hazard_types",
        joinColumns = @JoinColumn(name = "hazard_buffer_id"),
        inverseJoinColumns = @JoinColumn(name = "hazard_type_id"))
    private List<HazardType> restrictedHazardTypes = new ArrayList<>();
}""")

    add_h3(doc, "GisHazardBufferController.java — importGeoJson()")
    add_code(doc,
"""@PostMapping(value = "/import", consumes = MediaType.MULTIPART_FORM_DATA_VALUE)
@PreAuthorize("hasAuthority(...GIS_HAZARD_BUFFERS_MANAGE)")
public ResponseEntity<HazardBufferService.ImportResult> importGeoJson(
        @PathVariable String tenantSlug,
        @RequestParam("file") MultipartFile file,
        Authentication auth) {
    tenantAccessPolicy.assertTenantAccess(auth, tenantSlug);
    return ResponseEntity.ok(hazardBufferService.importGeoJson(file, tenantSlug));
}""")

    add_h3(doc, "ST_Intersects Query (GisHazardBufferRepository.java)")
    add_code(doc,
"""@Query(value = \"\"\"
    SELECT hb.* FROM gis_hazard_buffers hb
    JOIN gis_hazard_buffer_restricted_hazard_types rht
      ON rht.hazard_buffer_id = hb.id AND rht.hazard_type_id = :hazardTypeId
    WHERE ST_Intersects(hb.geometry,
        (SELECT geometry FROM gis_blocks
         WHERE layout_block_id = :layoutBlockId LIMIT 1))
    ORDER BY hb.name ASC, hb.id ASC
    \"\"\", nativeQuery = true)
List<GisHazardBuffer> findMatchingBuffersForLocation(
        @Param("layoutBlockId") UUID layoutBlockId,
        @Param("hazardTypeId") UUID hazardTypeId);""")

    doc.add_page_break()

    # ── Section 8 — Spatial Rule Enforcement ──────────────────────────────────
    add_h1(doc, "8 — Spatial Rule Enforcement")

    add_h2(doc, "Three-Step Validation Pipeline")
    embed_figure(doc, build_enforcement_diagram(), width_cm=10.0)

    add_body(doc,
        "Spatial rule enforcement is the operational core of the GIS subsystem. Every time a stock "
        "movement is requested — whether through a put-away operation or an inventory transfer — "
        "Spring Boot runs the three-step pipeline above before committing the movement. The "
        "orchestrator is InventoryValidationService, which holds the steps in strict order: "
        "hazard-buffer first, then zone category rules, then required-zone checks.")

    add_body(doc,
        "If a location has no GIS block on record, the service returns immediately without checking "
        "anything. This graceful-degrade behaviour means warehouses can be migrated incrementally: "
        "unmapped locations simply pass through without errors until an administrator draws their "
        "polygon on the map. Once a polygon exists, the full spatial pipeline applies.")

    add_h3(doc, "InventoryValidationService.java — assertLocationAllowsProduct()")
    add_code(doc,
"""@Transactional(readOnly = true)
public void assertLocationAllowsProduct(
        UUID locationId, Product product, boolean zoneOverride) {

    // No GIS block → no spatial context, skip all checks.
    if (gisBlockRepository.findByLayoutBlockId(locationId).isEmpty())
        return;

    // Step 1: Hazard-buffer check (always a hard BLOCK).
    assertNoHazardBufferViolation(locationId, product);

    // Step 2: Zone category-rule check (BLOCK or WARN depending on zone config).
    UUID categoryId = product.getCategory() != null
            ? product.getCategory().getId() : null;
    gisZoneValidationService.assertLocationAllowsProduct(
            locationId, categoryId, zoneOverride);

    // Step 3: Required-zone check (WARN only, reached only if steps 1-2 pass).
    assertRequiredZoneSatisfied(locationId, product);
}""")

    add_h3(doc, "Step 1 — Hazard Buffer Enforcement")
    add_code(doc,
"""private void assertNoHazardBufferViolation(UUID locationId, Product product) {
    HazardType hazardType = product.getHazardType();
    if (hazardType == null || "NONE".equals(hazardType.getCode()))
        return;

    List<GisHazardBuffer> matchingBuffers =
            geometryService.findMatchingHazardBuffers(locationId, hazardType.getId());

    if (!matchingBuffers.isEmpty()) {
        GisHazardBuffer first = matchingBuffers.get(0);
        List<HazardType> restricted =
                new ArrayList<>(first.getRestrictedHazardTypes());
        throw StorageRuleViolationException.hazardBufferBlock(first, restricted);
    }
}""")

    add_h3(doc, "Step 3 — Required Zone Type Enforcement")
    add_code(doc,
"""private void assertRequiredZoneSatisfied(UUID locationId, Product product) {
    ProductCategory category = product.getCategory();
    if (category == null) return;

    ZoneType required = category.getRequiredZoneType();
    if (required == null) return;

    boolean satisfied =
            geometryService.isLocationWithinZoneType(locationId, required.getId());
    if (!satisfied) {
        List<GisZone> suggested =
                geometryService.findZonesByZoneType(required.getId());
        throw StorageRuleViolationException.requiredZoneWarn(required, suggested);
    }
}""")

    add_h3(doc, "ST_Contains Query (GisZoneRepository.java)")
    add_code(doc,
"""// Find zones whose polygon spatially contains the location block:
@Query(value = \"\"\"
    SELECT z.* FROM gis_zones z
    WHERE ST_Contains(z.geometry,
        (SELECT geometry FROM gis_blocks
         WHERE layout_block_id = :layoutBlockId))
    \"\"\", nativeQuery = true)
List<GisZone> findZonesContainingLocation(
        @Param("layoutBlockId") UUID layoutBlockId);

// Check whether any zone of a required type contains the location:
@Query(value = \"\"\"
    SELECT COUNT(*) > 0 FROM gis_zones z
    WHERE z.zone_type_id = :zoneTypeId
      AND ST_Contains(z.geometry,
          (SELECT geometry FROM gis_blocks
           WHERE layout_block_id = :layoutBlockId LIMIT 1))
    \"\"\", nativeQuery = true)
boolean existsZoneOfTypeContainingLocation(
        @Param("layoutBlockId") UUID layoutBlockId,
        @Param("zoneTypeId") UUID zoneTypeId);""")

    doc.add_page_break()

    # ── Section 9 — Data Layers ────────────────────────────────────────────────
    add_h1(doc, "9 — Data Layers (Heatmap Overlays)")

    add_h2(doc, "Upload & Rendering Flow")
    embed_figure(doc, build_data_layers_diagram(), width_cm=15.5)

    add_body(doc,
        "Data layers are static image overlays that provide visual context on top of the warehouse "
        "map — density heatmaps, historical traffic patterns, or any other raster analysis produced "
        "in ArcGIS Pro and exported as a PNG or JPEG. They are purely presentational: they do not "
        "participate in validation logic and have no effect on stock movements. Their value is "
        "operational awareness, letting warehouse managers overlay analytical data on the live "
        "floor plan without switching tools.")

    add_body(doc,
        "Administrators upload image files through the web interface. Spring Boot validates the "
        "content type, stores the file on disk, and records the metadata in the gis_data_layers "
        "table. The image is served back to the browser as a raw byte array with the correct "
        "Content-Type header, which the ArcGIS SDK loads as a MediaLayer spanning the full "
        "warehouse extent. Users can toggle layer visibility, adjust opacity, and drag the image "
        "to fine-tune its position relative to the floor plan.")

    add_h3(doc, "GisDataLayerController.java — upload()")
    add_code(doc,
"""@PostMapping(consumes = MediaType.MULTIPART_FORM_DATA_VALUE)
@PreAuthorize("hasAuthority(...GIS_DATA_LAYERS_MANAGE)")
public ResponseEntity<GisDataLayerService.DataLayerSummary> upload(
        @PathVariable String tenantSlug,
        @RequestParam("name") String name,
        @RequestParam("file") MultipartFile file,
        Authentication auth) {
    tenantAccessPolicy.assertTenantAccess(auth, tenantSlug);
    return ResponseEntity.ok(dataLayerService.upload(tenantSlug, name, file));
}

@GetMapping("/{id}/image")
public ResponseEntity<byte[]> getImage(
        @PathVariable String tenantSlug,
        @PathVariable UUID id,
        Authentication auth) {
    tenantAccessPolicy.assertTenantAccess(auth, tenantSlug);
    ImageData imageData = dataLayerService.serveImage(id, tenantSlug);
    return ResponseEntity.ok()
            .contentType(imageData.mediaType())
            .body(imageData.bytes());
}""")

    doc.add_page_break()

    # ── Section 10 — Data Model Reference ─────────────────────────────────────
    add_h1(doc, "10 — Data Model Reference")

    add_body(doc,
        "The following table summarises the four PostGIS tables that underpin the GIS subsystem. "
        "All geometry columns use EPSG:4326 (WGS 84 geographic coordinates) to keep the data "
        "compatible with ArcGIS JS SDK 4, which defaults to that spatial reference.")

    add_data_model_table(doc)

    doc.add_page_break()

    # ── Section 11 — Spatial Query Reference ──────────────────────────────────
    add_h1(doc, "11 — Spatial Query Reference")

    add_body(doc,
        "Every geometry operation in the system flows through PostGIS. The table below lists the "
        "PostGIS functions in use and where each one appears in the Spring Boot layer.")

    add_spatial_query_table(doc)

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# Entry point
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    out_dir = os.path.join(os.path.dirname(__file__))
    out_path = os.path.join(out_dir, "warehouse_gis_technical_documentation.docx")

    print("Building document …")
    doc = build_document()
    doc.save(out_path)
    print(f"Saved → {out_path}")
